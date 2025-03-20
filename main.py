
#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
将 Markdown 文件完整转换为 Word (docx) 文件，并能正确识别以下类型的标记:
- 行内公式: $...$、\(...\)
- 块级公式: $$...$$、\[...\]
- Markdown 星号: *斜体*，**加粗**，***粗斜体***
- 表格单元格中的行内公式与星号
- 代码块(```...```)原样写入
- 图片引用: ![alt文本](image.png 或 http://xxx)，插入到 Word 并居中显示
  并且限制图片最大宽度为可用正文宽度的 80%。

在处理行内文本时，我们会统一拆分:
(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|\${2}.*?\${2}|\\\[.*?\\\]|\\\(.*?\)|\$.*?\$)

需要:
1) pip install latex2mathml python-docx lxml markdown requests
2) 在脚本同目录下放 MML2OMML.XSL (或自行修改 xsl_path).
"""

import os
import re
import requests
import markdown
import latex2mathml.converter
from latex2mathml.exceptions import NoAvailableTokensError
from lxml import etree
from io import BytesIO
from docx import Document
from docx import shared
from docx.shared import Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --------------------- 用于插入水平线 -----------------------------
def insert_horizontal_rule(doc):
    """
    在 doc 中插入一个“水平分割线”段落 (通过底部边框模拟)。
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert(0, pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')     # 线条样式
    bottom.set(qn('w:sz'), '6')          # 线条粗细
    bottom.set(qn('w:space'), '1')       # 与文字的间距
    bottom.set(qn('w:color'), 'auto')    # 颜色
    pBdr.append(bottom)
# --------------------- 新增部分结束 ----------------------------------------

def latex_to_omml(latex_input, xsl_path):
    """
    将单个 LaTeX 公式字符串转换为可供 python-docx 使用的 OMML XML 元素。
    带调试信息，查看传入与解析结果。
    """
    print(f"\n[DEBUG] -> latex_to_omml收到公式: {repr(latex_input)}")
    latex_input_stripped = latex_input.strip()
    if not latex_input_stripped:
        print("[DEBUG] -> 公式内容为空，返回 None")
        return None
    try:
        mathml_output = latex2mathml.converter.convert(latex_input_stripped)
        print(f"[DEBUG] -> latex2mathml转换成功, MathML片段(截断): {mathml_output[:80]}...")
        xsl_tree = etree.parse(xsl_path)
        transform = etree.XSLT(xsl_tree)
        mathml_dom = etree.fromstring(mathml_output)
        omml_dom = transform(mathml_dom)
        print(f"[DEBUG] -> XSLT转换成功, OMML(截断): {etree.tostring(omml_dom)[:80]}...")
        return omml_dom.getroot()
    except (NoAvailableTokensError, etree.XMLSyntaxError) as e:
        print(f"[DEBUG] -> latex2mathml或XML解析异常: {e}")
        return None
    except Exception as e:
        print(f"[DEBUG] -> 其他异常: {e}")
        return None

def strip_delimiters(latex_str):
    """
    去掉行内或块级分隔符 ($, $$, \(, \[ 等) 后得到纯 LaTeX 公式内容。
    带调试信息。
    """
    print(f" [DEBUG] strip_delimiters原始: {repr(latex_str)}")
    pattern = r'^(?:\${1,2}|\\\(|\\\[)\s*|\s*(?:\${1,2}|\\\)|\\\])$'
    result = re.sub(pattern, '', latex_str.strip())
    print(f" [DEBUG] strip_delimiters后: {repr(result)}")
    return result.strip()

def add_runs_with_inline_markdown(text, paragraph, xsl_path):
    """
    拆分行内文本, 同时支持:
    - Markdown星号: *...*, **...**, ***...***
    - LaTeX公式: $...$, $$...$$, \(...\), \[...\]
    并在python-docx的paragraph中插入对应的Run或OMML节点。
    """
    print(f"[DEBUG] -> 解析行内标记: {repr(text)}")
    inline_pattern = (
        r'(\*\*\*[^*]+\*\*\*|'    # ***粗斜体***
        r'\*\*[^*]+\*\*|'         # **加粗**
        r'\*[^*]+\*|'             # *斜体*
        r'\${2}[\s\S]*?\${2}|'    # $$公式$$
        r'\\\[[\s\S]*?\\\]|'      # \[公式\]
        r'\\\([\s\S]*?\\\)|'      # \(公式\)
        r'\$[\s\S]*?\$)'          # $公式$
    )
    segments = re.split(inline_pattern, text, flags=re.DOTALL)
    print(f"[DEBUG] -> re.split切分片段: {segments}")

    for seg in segments:
        # 如果 seg 能二次匹配说明它是一个公式或星号标记
        if re.match(inline_pattern, seg, flags=re.DOTALL):
            # (A) 判断是否是公式
            formula_pattern = r'(\${1,2}[\s\S]*?\${1,2}|\\\[[\s\S]*?\\\]|\\\([\s\S]*?\\\))'
            if re.match(formula_pattern, seg, flags=re.DOTALL):
                print(f"[DEBUG] -> 检测到公式段: {repr(seg)}")
                latex_code_inline = strip_delimiters(seg)
                omml_elem = latex_to_omml(latex_code_inline, xsl_path)
                if omml_elem is not None:
                    paragraph._element.append(omml_elem)
                else:
                    paragraph.add_run(seg)  # 转换失败则原样保留
            else:
                # (B) 星号标记
                print(f"[DEBUG] -> 检测到星号标记: {repr(seg)}")
                triple_star_pattern = r'^\*\*\*([^*]+)\*\*\*$'
                double_star_pattern = r'^\*\*([^*]+)\*\*$'
                single_star_pattern = r'^\*([^*]+)\*$'
                m3 = re.match(triple_star_pattern, seg)
                m2 = re.match(double_star_pattern, seg)
                m1 = re.match(single_star_pattern, seg)
                if m3:
                    content = m3.group(1)
                    run_obj = paragraph.add_run(content)
                    run_obj.bold = True
                    run_obj.italic = True
                elif m2:
                    content = m2.group(1)
                    run_obj = paragraph.add_run(content)
                    run_obj.bold = True
                elif m1:
                    content = m1.group(1)
                    run_obj = paragraph.add_run(content)
                    run_obj.italic = True
                else:
                    paragraph.add_run(seg)
        else:
            # 普通文本
            paragraph.add_run(seg)

def get_usable_width(doc):
    """
    计算Word文档的内容区宽度(单位: EMU) = 页面宽度 - 左右边距。
    """
    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    return page_width - left_margin - right_margin

def handle_image(doc, alt_text, image_url, md_file_path):
    """
    将图片插入 Word：
    - 若 image_url 是网络地址，则下载后插入；
    - 若是本地(相对/绝对)路径，则直接插入；
    - 将所在段落居中对齐，并限制图片不超过内容区宽度的80%(高度自适应)。
    """
    print(f"[DEBUG] -> 插入图片, alt_text={alt_text}, url={image_url}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    # 可用宽度(EMU)
    max_width_emu = int(get_usable_width(doc) * 0.8)

    if image_url.lower().startswith('http://') or image_url.lower().startswith('https://'):
        # 网络图片
        try:
            resp = requests.get(image_url, timeout=10)
            if resp.status_code == 200:
                image_data = BytesIO(resp.content)
                run.add_picture(image_data, width=Emu(max_width_emu))
            else:
                doc.add_paragraph(f"(图片加载失败: {image_url}, status={resp.status_code})")
        except Exception as e:
            doc.add_paragraph(f"(下载图片时出错: {image_url}, err={e})")
    else:
        # 本地图片
        if not os.path.isabs(image_url):
            md_dir = os.path.dirname(os.path.abspath(md_file_path))
            image_url = os.path.join(md_dir, image_url)
        if os.path.exists(image_url):
            try:
                run.add_picture(image_url, width=Emu(max_width_emu))
            except Exception as e:
                doc.add_paragraph(f"(插入图片时出错: {image_url}, err={e})")
        else:
            doc.add_paragraph(f"(图片未找到: {image_url})")

def parse_line_for_images_and_text(line, doc, xsl_path, md_file_path):
    """
    识别行内的 Markdown 图片引用(![alt](url)) 并插入 Word(含居中、宽度限制)，
    剩余文本则交给 add_runs_with_inline_markdown() 处理斜体/加粗/公式等。
    """
    img_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'

    segments = []
    last_idx = 0

    for match in re.finditer(img_pattern, line):
        start, end = match.span()
        # 图片前的普通文本
        if start > last_idx:
            segments.append(("text", line[last_idx:start]))
        alt_text = match.group(1)
        img_url = match.group(2)
        segments.append(("image", alt_text, img_url))
        last_idx = end

    # 结尾剩余文本
    if last_idx < len(line):
        segments.append(("text", line[last_idx:]))

    # 依次写入
    for seg in segments:
        if seg[0] == "text":
            text_content = seg[1]
            # 如果只有空白，留一个空段落
            if text_content.strip():
                paragraph = doc.add_paragraph(style=None)
                add_runs_with_inline_markdown(text_content, paragraph, xsl_path)
            else:
                doc.add_paragraph("")
        else:
            # seg[0] == "image"
            alt_text, img_url = seg[1], seg[2]
            handle_image(doc, alt_text, img_url, md_file_path)

def convert_markdown_to_docx(md_file, xsl_path, output_docx):
    """
    主函数：将 Markdown 文件转换为 Word (docx).
    - 解析标题、表格、公式、代码块、图片、列表等
    - 自动将图片限制为内容区宽度的80%，并居中
    - 最终保存为 output_docx
    """
    print(f"[DEBUG] -> 开始读取Markdown文件: {md_file}")
    with open(md_file, 'r', encoding='utf-8') as f:
        md_text = f.read()
    print("[DEBUG] -> 读取完成, 字符数: ", len(md_text))

    # 测试 markdown -> html (不做实际使用)，仅验证 markdown 包
    _ = markdown.markdown(md_text, extensions=["tables", "fenced_code"])

    doc = Document()
    # 设置全局的字体/字号
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(10)

    lines = md_text.split('\n')

    in_code_block = False
    in_math_block = False
    math_block_buffer = []

    # 表格相关
    table_buffer = []
    table_mode = False

    def flush_table(doc_obj):
        """
        将 table_buffer 中的行转换成表格插入 docObj，并清空。
        表格单元格中也支持行内公式、星号标记，但暂不处理图片。
        """
        nonlocal table_buffer
        if not table_buffer:
            return

        print(f"[DEBUG] -> flush_table, 行数: {len(table_buffer)}")
        rows_data = []
        for row in table_buffer:
            row = row.strip().strip('|')
            row_cells = [cell.strip() for cell in row.split('|')]
            rows_data.append(row_cells)

        nrows = len(rows_data)
        ncols = len(rows_data[0]) if rows_data else 0
        if nrows == 0 or ncols == 0:
            table_buffer = []
            return

        table = doc_obj.add_table(rows=nrows, cols=ncols)
        table.style = 'Table Grid'
        for i in range(nrows):
            for j in range(ncols):
                cell_paragraph = table.cell(i, j).paragraphs[0]
                cell_paragraph.text = ''
                cell_text = rows_data[i][j]
                add_runs_with_inline_markdown(cell_text, cell_paragraph, xsl_path)

        table_buffer = []

    print("[DEBUG] -> 开始逐行处理Markdown...")
    for line in lines:
        line_stripped = line.strip()
        print(f"\n[DEBUG] -> 处理行: {repr(line_stripped)}")

        # 1. 代码块检测
        if line_stripped.startswith('```'):
            in_code_block = not in_code_block
            print(f"[DEBUG] -> {'进入' if in_code_block else '退出'}代码块")
            if in_code_block:
                doc.add_paragraph("----- 代码块开始 -----")
            else:
                doc.add_paragraph("----- 代码块结束 -----")
            continue

        if in_code_block:
            # 代码块区域，原样写入
            doc.add_paragraph(line)
            continue

        # 2. 块公式是否在进行中
        if in_math_block:
            if line_stripped.endswith('$$') or line_stripped.endswith('\\]'):
                math_block_buffer.append(line)
                print("[DEBUG] -> 块公式结束, 开始转换...")
                in_math_block = False
                block_content = "\n".join(math_block_buffer)
                block_content_stripped = strip_delimiters(block_content)
                paragraph = doc.add_paragraph()
                omml_elem = latex_to_omml(block_content_stripped, xsl_path)
                if omml_elem is not None:
                    paragraph._element.append(omml_elem)
                else:
                    paragraph.add_run(block_content)
                math_block_buffer = []
                continue
            else:
                math_block_buffer.append(line)
                continue

        # 3. 检查多行块公式的开始
        if (
            (line_stripped.startswith('$$') or line_stripped.startswith('\\[')) and
            not (line_stripped.endswith('$$') or line_stripped.endswith('\\]'))
        ):
            print("[DEBUG] -> 多行块公式开始")
            in_math_block = True
            math_block_buffer = [line]
            continue

        # 4. 单行块公式
        block_patt = r'^(?:\${2}.*?\${2}|\\\[.*?\\\])$'
        if re.match(block_patt, line_stripped):
            print("[DEBUG] -> 单行块公式")
            latex_code_block = strip_delimiters(line_stripped)
            paragraph = doc.add_paragraph()
            omml_elem = latex_to_omml(latex_code_block, xsl_path)
            if omml_elem is not None:
                paragraph._element.append(omml_elem)
            else:
                paragraph.add_run(line_stripped)
            continue

        # 5. 空行
        if not line_stripped:
            print("[DEBUG] -> 空行")
            if table_mode:
                flush_table(doc)
                table_mode = False
            doc.add_paragraph("")
            continue

        # 6. 表格行
        if line_stripped.startswith('|') and line_stripped.endswith('|'):
            # 如果整行只含 '-', '|', ':' 等字符(典型markdown表头分隔)，跳过
            if all(ch in '-|: ' for ch in line_stripped):
                print("[DEBUG] -> 检测到 Markdown 表格分隔行，跳过")
                continue
            print("[DEBUG] -> 表格行")
            table_buffer.append(line_stripped)
            table_mode = True
            continue
        else:
            if table_mode:
                flush_table(doc)
                table_mode = False

        # 7. 水平线 "---"
        if line_stripped == '---':
            print("[DEBUG] -> 检测到水平线标记 '---'")
            insert_horizontal_rule(doc)
            continue

        # 8. 标题(# 开头)
        if line_stripped.startswith('#'):
            match_hashes = re.match(r'^(#+)', line_stripped)
            level = len(match_hashes.group(1)) if match_hashes else 1
            heading_text = line_stripped[level:].strip()
            print(f"[DEBUG] -> 标题: level={level}, text={heading_text}")
            doc.add_heading(heading_text, min(level, 6))
            continue

        # 9. 无序列表
        if line_stripped.startswith('* ') or line_stripped.startswith('- '):
            item_text = line_stripped[1:].strip()
            print(f"[DEBUG] -> 无序列表项: {item_text}")
            paragraph = doc.add_paragraph(style=None)
            paragraph.add_run("• " + item_text)
            continue

        # 10. 普通文本行(可能含行内公式、星号或图片)
        print("[DEBUG] -> 普通文本行(可能含图片/公式/星号)")
        parse_line_for_images_and_text(line, doc, xsl_path, md_file)

    # 文档末尾，可能还有表格缓存
    if table_mode:
        flush_table(doc)

    doc.save(output_docx)
    print(f"[DEBUG] -> 已生成 Word 文档: {output_docx}")

if __name__ == "__main__":
    # 设置 XSLT 路径
    xsl_path = os.path.join('.', 'MML2OMML.XSL')
    md_file = "example.md"  # 你的Markdown文件路径
    output_docx = "79utput_with_scaled_images.docx"
    convert_markdown_to_docx(md_file, xsl_path, output_docx)
